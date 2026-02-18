"""
Enhanced DOCX Processor.
Extracts heading hierarchy, headers/footers, footnotes, merged cell tables, hyperlinks.
"""

import io
import os
import re
import logging
import tempfile
import subprocess
from typing import List, Dict, Any

from processors.base import FileProcessor, TextExtractionResult

logger = logging.getLogger(__name__)


class DOCXProcessor(FileProcessor):
    """Process DOCX files with enhanced structural extraction."""

    def can_process(self, mime_type: str, extension: str) -> bool:
        return (mime_type in [
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/msword'
        ] or extension.lower() in ['.docx', '.doc'])

    def extract_text(self, file_bytes: bytes, filename: str, file_id: str) -> TextExtractionResult:
        # For .doc files, try pandoc fallback
        if filename.lower().endswith('.doc'):
            try:
                return self._extract_with_pandoc(file_bytes, filename, file_id)
            except Exception as e:
                logger.warning(f"Pandoc fallback failed for .doc: {e}")
                raise

        result = TextExtractionResult(file_id, filename, "docx")
        result.extraction_method = "python-docx-enhanced"

        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(io.BytesIO(file_bytes))
            full_text = []

            # Extract headers and footers from all sections
            header_texts, footer_texts = self._extract_headers_footers(doc)
            if header_texts:
                full_text.append(f"[HEADER]: {' | '.join(header_texts)}")

            # Extract paragraphs with heading hierarchy
            for para in doc.paragraphs:
                if not para.text.strip():
                    continue

                prefix = ""
                if para.style and para.style.name:
                    style_name = para.style.name
                    if style_name.startswith("Heading"):
                        level_str = style_name.replace("Heading ", "").replace("Heading", "1")
                        try:
                            level = int(level_str)
                        except ValueError:
                            level = 1
                        prefix = "#" * level + " "

                full_text.append(prefix + para.text)

            # Extract hyperlinks
            hyperlinks = self._extract_hyperlinks(doc)

            # Extract tables with merged cell handling
            table_count = 0
            for table in doc.tables:
                table_data = self._extract_table_with_merged_cells(table)
                if table_data:
                    result.add_table(1, table_data, table_count)
                    table_text = "\n[TABLE]\n" + "\n".join(" | ".join(row) for row in table_data)
                    full_text.append(table_text)
                    table_count += 1

            # Extract footnotes
            footnotes = self._extract_footnotes(doc)
            if footnotes:
                full_text.append("\n[FOOTNOTES]")
                full_text.extend(footnotes)

            # Extract images info
            images = self._extract_images(doc)
            for img_idx, img_info in enumerate(images):
                result.add_image(1, img_idx, img_info)

            if footer_texts:
                full_text.append(f"[FOOTER]: {' | '.join(footer_texts)}")

            complete_text = '\n'.join(full_text)
            result.add_page(1, complete_text, "python-docx-enhanced")

            result.quality_metrics.average_confidence = 100.0
            result.quality_metrics.completeness_score = min(100, len(complete_text) / 100)

            result.metadata = {
                "paragraph_count": len(doc.paragraphs),
                "table_count": table_count,
                "image_count": len(images),
                "hyperlink_count": len(hyperlinks),
                "footnote_count": len(footnotes),
                "has_headers": bool(header_texts),
                "has_footers": bool(footer_texts),
                "total_characters": len(complete_text)
            }

            logger.info(
                f"Extracted text from DOCX: {filename} "
                f"({table_count} tables, {len(images)} images, {len(footnotes)} footnotes)"
            )

        except Exception as e:
            logger.error(f"Failed to process DOCX {filename}: {e}")
            try:
                return self._extract_with_pandoc(file_bytes, filename, file_id)
            except Exception:
                raise e

        return result

    def _extract_headers_footers(self, doc) -> tuple:
        """Extract headers and footers from all sections."""
        headers = []
        footers = []
        try:
            for section in doc.sections:
                if section.header:
                    header_text = " ".join(
                        p.text for p in section.header.paragraphs if p.text.strip()
                    )
                    if header_text.strip() and header_text not in headers:
                        headers.append(header_text.strip())

                if section.footer:
                    footer_text = " ".join(
                        p.text for p in section.footer.paragraphs if p.text.strip()
                    )
                    if footer_text.strip() and footer_text not in footers:
                        footers.append(footer_text.strip())
        except Exception as e:
            logger.debug(f"Header/footer extraction failed: {e}")
        return headers, footers

    def _extract_hyperlinks(self, doc) -> List[Dict[str, str]]:
        """Extract hyperlinks from the document."""
        hyperlinks = []
        try:
            for rel in doc.part.rels.values():
                if "hyperlink" in rel.reltype:
                    hyperlinks.append({
                        "target": rel.target_ref,
                        "rel_id": rel.rId
                    })
        except Exception as e:
            logger.debug(f"Hyperlink extraction failed: {e}")
        return hyperlinks

    def _extract_table_with_merged_cells(self, table) -> List[List[str]]:
        """Extract table data handling merged cells (python-docx repeats merged cell text)."""
        rows = []
        for row in table.rows:
            row_data = []
            prev_text = None
            for cell in row.cells:
                text = cell.text.strip()
                # python-docx repeats the same Cell object for merged cells
                # Detect by checking if it's the exact same object or text as previous
                if prev_text is not None and text == prev_text and id(cell) == id(row.cells[len(row_data) - 1] if row_data else None):
                    row_data.append("")  # Mark merged cell as empty to avoid duplication
                else:
                    row_data.append(text)
                prev_text = text
            rows.append(row_data)
        return rows

    def _extract_footnotes(self, doc) -> List[str]:
        """Extract footnotes by parsing the document's XML package."""
        footnotes = []
        try:
            from lxml import etree
            # Access footnotes part from the document package
            for rel in doc.part.rels.values():
                if 'footnotes' in rel.target_ref.lower():
                    footnotes_part = rel.target_part
                    root = etree.fromstring(footnotes_part.blob)
                    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                    for footnote in root.findall('.//w:footnote', ns):
                        fn_id = footnote.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id', '')
                        if fn_id in ('0', '-1'):  # Skip separator and continuation
                            continue
                        texts = []
                        for t_elem in footnote.findall('.//w:t', ns):
                            if t_elem.text:
                                texts.append(t_elem.text)
                        if texts:
                            footnotes.append(f"[{fn_id}] {''.join(texts)}")
                    break
        except Exception as e:
            logger.debug(f"Footnote extraction failed: {e}")
        return footnotes

    def _extract_images(self, doc) -> List[Dict[str, Any]]:
        """Extract image metadata from DOCX."""
        images = []
        try:
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    images.append({
                        "type": "embedded_image",
                        "format": rel.target_ref.split('.')[-1] if '.' in rel.target_ref else "unknown",
                        "relationship_id": rel.rId
                    })
        except Exception as e:
            logger.debug(f"Image extraction from DOCX failed: {e}")
        return images

    def _extract_with_pandoc(self, file_bytes: bytes, filename: str, file_id: str) -> TextExtractionResult:
        """Fallback extraction using pandoc (supports .doc and .docx)."""
        result = TextExtractionResult(file_id, filename, "docx")
        result.extraction_method = "pandoc"

        suffix = '.doc' if filename.lower().endswith('.doc') else '.docx'
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as temp_file:
            temp_file.write(file_bytes)
            temp_file.flush()

            try:
                output = subprocess.check_output(
                    ['pandoc', temp_file.name, '-t', 'plain', '--wrap=none'],
                    stderr=subprocess.STDOUT,
                    timeout=60
                )
                text = output.decode('utf-8')
                result.add_page(1, text, "pandoc")
                result.quality_metrics.average_confidence = 80.0
                result.quality_metrics.completeness_score = min(100, len(text) / 100)
            finally:
                os.unlink(temp_file.name)

        return result
